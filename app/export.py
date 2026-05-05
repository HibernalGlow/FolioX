# -*- coding: utf-8 -*-
"""DOCX export — parse OCR text and build Word documents."""
import io
import re

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel

from .ocr.postprocess import TableParser


class ExportPage(BaseModel):
    num: int
    text: str


class ExportRequest(BaseModel):
    pages: list[ExportPage]
    title: str | None = None


def parse_md_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        stripped = line.strip()
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if all(re.match(r'^[-:]+$', c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


def parse_ocr_text(text: str) -> list[dict]:
    """Parse OCR text into structured elements: headings, paragraphs, tables."""
    elements = []

    parts = re.split(r'(<table[\s\S]*?</table>)', text, flags=re.IGNORECASE)

    for part in parts:
        if re.match(r'^<table[\s\S]*</table>$', part, re.IGNORECASE):
            parser = TableParser()
            parser.feed(part)
            if parser.rows:
                elements.append({
                    "type": "table",
                    "rows": parser.rows,
                    "header_rows": parser.header_row_indices,
                })
            continue

        lines = part.split('\n')
        md_table_buf = []

        def flush_md_table():
            if not md_table_buf:
                return
            rows = parse_md_table(md_table_buf)
            if rows:
                elements.append({
                    "type": "table",
                    "rows": rows,
                    "header_rows": {0},
                })
            md_table_buf.clear()

        for line in lines:
            trimmed = line.strip()

            if '|' in trimmed and (trimmed.startswith('|') or re.search(r'\w\s*\|', trimmed)):
                md_table_buf.append(trimmed)
                continue

            if md_table_buf:
                flush_md_table()

            if trimmed.startswith('### '):
                elements.append({"type": "heading", "level": 3, "text": trimmed[4:]})
            elif trimmed.startswith('## '):
                elements.append({"type": "heading", "level": 2, "text": trimmed[3:]})
            elif trimmed.startswith('# '):
                elements.append({"type": "heading", "level": 1, "text": trimmed[2:]})
            elif trimmed:
                elements.append({"type": "paragraph", "text": trimmed})

        flush_md_table()

    return elements


def build_docx(title: str, pages: list[ExportPage]) -> io.BytesIO:
    """Build a real DOCX file from parsed OCR text."""
    doc = DocxDocument()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    ea_font = rpr.makeelement(qn('w:rFonts'), {qn('w:eastAsia'): '微软雅黑'})
    rpr.append(ea_font)

    multi_page = len(pages) > 1

    if multi_page:
        doc.add_heading(title, level=0)

    for idx, page in enumerate(pages):
        if multi_page and idx > 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)

        elements = parse_ocr_text(page.text or '')

        for elem in elements:
            if elem["type"] == "heading":
                doc.add_heading(elem["text"], level=elem["level"])

            elif elem["type"] == "paragraph":
                para = doc.add_paragraph()
                text = elem["text"]
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
                for p in parts:
                    if p.startswith('**') and p.endswith('**'):
                        run = para.add_run(p[2:-2])
                        run.bold = True
                    elif p.startswith('*') and p.endswith('*') and len(p) > 2:
                        run = para.add_run(p[1:-1])
                        run.italic = True
                    else:
                        para.add_run(p)

            elif elem["type"] == "table":
                rows = elem["rows"]
                if not rows:
                    continue
                n_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=n_cols)
                tbl.style = 'Table Grid'

                header_rows = elem.get("header_rows", set())

                for i, row_data in enumerate(rows):
                    row = tbl.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < n_cols:
                            cell = row.cells[j]
                            cell.text = cell_text
                            if i in header_rows:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.bold = True

    if multi_page:
        for idx, section in enumerate(doc.sections):
            footer = section.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0]
            para.text = f"— {pages[idx].num} —"
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
